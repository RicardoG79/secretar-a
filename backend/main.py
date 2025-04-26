import pdfplumber
import re
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from datetime import datetime
import sqlite3
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import io
import os

app = FastAPI()

# Habilitar CORS para comunicación con el frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Inicializar base de datos SQLite
def init_db():
    conn = sqlite3.connect("database.db")
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS resultados
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  rol_usuario TEXT,
                  nombre_archivo TEXT,
                  dias_totales INTEGER,
                  ruta_reporte TEXT,
                  fecha_creacion TEXT)''')
    conn.commit()
    conn.close()

init_db()

# Función para parsear fechas en formato dd/mm/aa
def parsear_fecha(fecha_str):
    try:
        return datetime.strptime(fecha_str, "%d/%m/%y")
    except ValueError:
        return None

# Función para calcular días trabajados
def calcular_dias_trabajados(periodos, fecha_fin_str):
    fecha_fin = parsear_fecha(fecha_fin_str)
    if not fecha_fin:
        fecha_fin = datetime.now()
    
    dias_totales = 0
    for periodo in periodos:
        fecha_inicio = parsear_fecha(periodo["inicio"])
        fecha_fin_periodo = parsear_fecha(periodo["fin"]) if periodo["fin"] else fecha_fin
        if fecha_inicio and fecha_fin_periodo:
            dias = (fecha_fin_periodo - fecha_inicio).days
            if dias > 0:
                dias_totales += dias
    return dias_totales

# Función para extraer periodos del PDF
def extraer_periodos_de_pdf(contenido_archivo):
    periodos = []
    with pdfplumber.open(io.BytesIO(contenido_archivo)) as pdf:
        texto = ""
        for pagina in pdf.pages:
            texto += pagina.extract_text() or ""
        
        # Extraer fecha de emisión (ej., "veinticinco dias del mes de abril del año dos mil veinticinco")
        coincidencia_emision = re.search(r"a los (\w+) dias del mes de (\w+) del año dos mil (\w+)", texto, re.IGNORECASE)
        fecha_emision = None
        if coincidencia_emision:
            dia = coincidencia_emision.group(1)
            mes = coincidencia_emision.group(2)
            anio = coincidencia_emision.group(3)
            # Mapear texto en español a números
            mapa_dias = {"veinticinco": "25"}
            mapa_meses = {"abril": "04"}
            mapa_anios = {"veinticinco": "25"}
            fecha_emision = f"{mapa_dias.get(dia, '01')}/{mapa_meses.get(mes, '01')}/{mapa_anios.get(anio, '25')}"
        
        # Extraer filas de la tabla
        lineas = texto.split("\n")
        for linea in lineas:
            coincidencia = re.match(r"(.+?)\s+\d+\s+(Suplente|Titular)\s+\d+/\d+\s+(\d{2})\s+(\d{2})\s+(\d{2})\s+(\d{2})?\s*(\d{2})?\s*(\d{2})?", linea)
            if coincidencia:
                dia_inicio, mes_inicio, anio_inicio = coincidencia.group(3), coincidencia.group(4), coincidencia.group(5)
                dia_fin, mes_fin, anio_fin = coincidencia.group(6), coincidencia.group(7), coincidencia.group(8)
                fecha_inicio = f"{dia_inicio}/{mes_inicio}/{anio_inicio}"
                fecha_fin = f"{dia_fin}/{mes_fin}/{anio_fin}" if dia_fin and mes_fin and anio_fin else ""
                periodos.append({"inicio": fecha_inicio, "fin": fecha_fin})
    
    return periodos, fecha_emision

# Endpoint para subir y procesar PDF
@app.post("/subir/")
async def subir_archivo(archivo: UploadFile = File(...), rol_usuario: str = "usuario1", fecha_fin: str = None):
    if not archivo.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Solo se permiten archivos PDF")
    
    contenido = await archivo.read()
    periodos, fecha_emision = extraer_periodos_de_pdf(contenido)
    
    # Usar fecha_fin proporcionada o fecha de emisión
    fecha_final = fecha_fin or fecha_emision or datetime.now().strftime("%d/%m/%y")
    dias_totales = calcular_dias_trabajados(periodos, fecha_final)
    
    # Generar reporte PDF
    buffer_pdf = io.BytesIO()
    c = canvas.Canvas(buffer_pdf, pagesize=letter)
    c.drawString(100, 750, f"Reporte de Tiempo Trabajado para {rol_usuario}")
    c.drawString(100, 730, f"Días Totales Trabajados: {dias_totales}")
    c.drawString(100, 710, f"Calculado Hasta: {fecha_final}")
    c.save()
    buffer_pdf.seek(0)
    
    # Generar reporte Word
    doc = Document()
    doc.add_heading(f"Reporte de Tiempo Trabajado para {rol_usuario}", 0)
    doc.add_paragraph(f"Días Totales Trabajados: {dias_totales}")
    doc.add_paragraph(f"Calculado Hasta: {fecha_final}")
    buffer_word = io.BytesIO()
    doc.save(buffer_word)
    buffer_word.seek(0)
    
    # Guardar en base de datos
    conn = sqlite3.connect("database.db")
    c = conn.cursor()
    c.execute("INSERT INTO resultados (rol_usuario, nombre_archivo, dias_totales, ruta_reporte, fecha_creacion) VALUES (?, ?, ?, ?, ?)",
              (rol_usuario, archivo.filename, dias_totales, f"reportes/{archivo.filename}.pdf", datetime.now().isoformat()))
    conn.commit()
    conn.close()
    
    # Guardar reportes en disco
    os.makedirs("reportes", exist_ok=True)
    with open(f"reportes/{archivo.filename}.pdf", "wb") as f:
        f.write(buffer_pdf.getvalue())
    with open(f"reportes/{archivo.filename}.docx", "wb") as f:
        f.write(buffer_word.getvalue())
    
    return {
        "dias_totales": dias_totales,
        "url_pdf": f"/reportes/{archivo.filename}.pdf",
        "url_word": f"/reportes/{archivo.filename}.docx",
        "fecha_fin": fecha_final
    }

# Endpoint para servir reportes
from fastapi.staticfiles import StaticFiles
app.mount("/reportes", StaticFiles(directory="reportes"), name="reportes")